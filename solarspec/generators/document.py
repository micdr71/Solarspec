"""Technical specification document generator."""

from __future__ import annotations

import html as html_module
from datetime import date

from solarspec.models import SystemDesign


def _escape_html(text: str) -> str:
    """Escape text for safe HTML insertion, preserving newlines as <br>."""
    return html_module.escape(text).replace("\n", "<br>")


def generate(
    design: SystemDesign,
    output_path: str,
    format: str = "docx",
    narrative: dict[str, str] | None = None,
) -> str:
    """Generate a technical specification document.

    Args:
        design: Complete system design.
        output_path: Output file path.
        format: Output format ('docx' or 'pdf').
        narrative: Optional AI-generated narrative sections dict.

    Returns:
        Path to the generated document.
    """
    if format == "docx":
        return _generate_docx(design, output_path, narrative=narrative)
    elif format == "pdf":
        return _generate_pdf(design, output_path, narrative=narrative)
    else:
        raise ValueError(f"Formato non supportato: {format}. Usa 'docx' o 'pdf'.")


def _generate_docx(
    design: SystemDesign, output_path: str, narrative: dict[str, str] | None = None
) -> str:
    """Generate a DOCX technical specification."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise ImportError("Installa python-docx: pip install python-docx")

    narr = narrative or {}
    doc = Document()

    # Title
    doc.add_heading("Capitolato Tecnico — Impianto Fotovoltaico", level=0)

    # Premessa (AI narrative)
    if narr.get("premessa"):
        doc.add_heading("Premessa", level=1)
        doc.add_paragraph(narr["premessa"])

    # Site info
    doc.add_heading("1. Dati del sito", level=1)
    if narr.get("analisi_sito"):
        doc.add_paragraph(narr["analisi_sito"])
    doc.add_paragraph(f"Indirizzo: {design.site.address}")
    doc.add_paragraph(
        f"Coordinate: {design.site.latitude:.5f}°N, {design.site.longitude:.5f}°E"
    )
    doc.add_paragraph(f"Comune: {design.site.municipality} ({design.site.province})")
    doc.add_paragraph(f"Zona climatica: {design.site.climate_zone}")
    doc.add_paragraph(f"Zona sismica: {design.site.seismic_zone}")

    # Solar data
    doc.add_heading("2. Analisi solare", level=1)
    if narr.get("risorsa_solare"):
        doc.add_paragraph(narr["risorsa_solare"])
    doc.add_paragraph(
        f"Irraggiamento annuo (piano ottimale): {design.solar_data.annual_irradiation} kWh/m²/anno"
    )
    doc.add_paragraph(f"Inclinazione ottimale: {design.solar_data.optimal_tilt}°")
    doc.add_paragraph(f"Azimut ottimale: {design.solar_data.optimal_azimuth}°")
    doc.add_paragraph(
        f"Producibilità specifica: {design.solar_data.annual_production_per_kwp} kWh/kWp/anno"
    )

    # System design
    doc.add_heading("3. Dimensionamento impianto", level=1)
    if narr.get("dimensionamento"):
        doc.add_paragraph(narr["dimensionamento"])
    doc.add_paragraph(f"Potenza nominale: {design.system_size_kwp} kWp")
    doc.add_paragraph(f"Numero moduli: {design.num_panels}")
    if design.module:
        doc.add_paragraph(
            f"Modulo: {design.module.manufacturer} {design.module.model} "
            f"({design.module.power_wp} Wp, η={design.module.efficiency}%)"
        )
    if design.inverter:
        doc.add_paragraph(
            f"Inverter: {design.inverter.manufacturer} {design.inverter.model} "
            f"({design.inverter.power_kw} kW, η={design.inverter.efficiency}%)"
        )
    doc.add_paragraph(f"Produzione annua stimata: {design.estimated_production_kwh:.0f} kWh")
    doc.add_paragraph(f"Autoconsumo stimato: {design.self_consumption_rate}%")
    doc.add_paragraph(f"Performance Ratio: {design.performance_ratio}")

    # Economics
    if design.economics:
        doc.add_heading("4. Analisi economica", level=1)
        if narr.get("analisi_economica"):
            doc.add_paragraph(narr["analisi_economica"])
        doc.add_paragraph(f"Costo totale stimato: €{design.economics.total_cost_eur:,.2f}")
        doc.add_paragraph(f"Costo per kWp: €{design.economics.cost_per_kwp:,.2f}/kWp")
        doc.add_paragraph(f"Risparmio annuo stimato: €{design.economics.annual_savings_eur:,.2f}")
        doc.add_paragraph(f"Tempo di rientro: {design.economics.payback_years} anni")
        doc.add_paragraph(f"ROI a 25 anni: {design.economics.roi_25y_percent}%")
        doc.add_paragraph(f"LCOE: €{design.economics.lcoe}/kWh")
        doc.add_paragraph(f"Incentivo: {design.economics.incentive_type}")

    # Notes
    if design.notes:
        doc.add_heading("5. Note", level=1)
        for note in design.notes:
            doc.add_paragraph(f"• {note}")

    # Normativa
    doc.add_heading("6. Riferimenti normativi", level=1)
    norms = [
        "CEI 0-21 — Regola tecnica di connessione utenti attivi BT",
        "CEI 0-16 — Regola tecnica di connessione utenti attivi MT",
        "D.Lgs. 199/2021 — Attuazione direttiva RED II",
        "DM 14/01/2008 — Norme tecniche costruzioni (NTC)",
    ]
    for norm in norms:
        doc.add_paragraph(f"• {norm}")

    # Conclusioni (AI narrative)
    if narr.get("conclusioni"):
        doc.add_heading("7. Conclusioni e raccomandazioni", level=1)
        doc.add_paragraph(narr["conclusioni"])

    # Footer
    doc.add_paragraph("")
    doc.add_paragraph("Documento generato con SolarSpec — https://github.com/micdr71/Solarspec")

    doc.save(output_path)
    return output_path


def _build_html(design: SystemDesign, narrative: dict[str, str] | None = None) -> str:
    """Build an HTML representation of the technical specification."""
    narr = narrative or {}
    today = date.today().strftime("%d/%m/%Y")
    monthly_labels = [
        "Gen", "Feb", "Mar", "Apr", "Mag", "Giu",
        "Lug", "Ago", "Set", "Ott", "Nov", "Dic",
    ]
    monthly_data = design.solar_data.monthly_irradiation or []

    module_html = ""
    if design.module:
        module_html = (
            f"<tr><td>Modulo</td><td>{design.module.manufacturer} {design.module.model} "
            f"({design.module.power_wp} Wp, &eta;={design.module.efficiency}%)</td></tr>"
        )

    inverter_html = ""
    if design.inverter:
        inverter_html = (
            f"<tr><td>Inverter</td><td>{design.inverter.manufacturer} {design.inverter.model} "
            f"({design.inverter.power_kw} kW, &eta;={design.inverter.efficiency}%)</td></tr>"
        )

    analisi_econ_narr = ""
    if narr.get("analisi_economica"):
        analisi_econ_narr = f'<p class="narrative">{_escape_html(narr["analisi_economica"])}</p>'

    economics_html = ""
    if design.economics:
        economics_html = f"""
        <h2>4. Analisi economica</h2>
        {analisi_econ_narr}
        <table>
            <tr><td>Costo totale stimato</td><td>&euro;{design.economics.total_cost_eur:,.2f}</td></tr>
            <tr><td>Costo per kWp</td><td>&euro;{design.economics.cost_per_kwp:,.2f}/kWp</td></tr>
            <tr><td>Risparmio annuo stimato</td><td>&euro;{design.economics.annual_savings_eur:,.2f}</td></tr>
            <tr><td>Tempo di rientro</td><td>{design.economics.payback_years} anni</td></tr>
            <tr><td>ROI a 25 anni</td><td>{design.economics.roi_25y_percent}%</td></tr>
            <tr><td>LCOE</td><td>&euro;{design.economics.lcoe}/kWh</td></tr>
            <tr><td>Incentivo</td><td>{design.economics.incentive_type}</td></tr>
            <tr><td>Valore incentivi (25 anni)</td><td>&euro;{design.economics.incentive_value_eur:,.2f}</td></tr>
        </table>
        """

    notes_html = ""
    if design.notes:
        items = "".join(f"<li>{n}</li>" for n in design.notes)
        notes_html = f"<h2>5. Note</h2><ul>{items}</ul>"

    monthly_rows = ""
    for i, val in enumerate(monthly_data):
        label = monthly_labels[i] if i < len(monthly_labels) else str(i + 1)
        monthly_rows += f"<tr><td>{label}</td><td>{val} kWh/m&sup2;</td></tr>"

    premessa_html = ""
    if narr.get("premessa"):
        premessa_html = f'<h2>Premessa</h2><p class="narrative">{_escape_html(narr["premessa"])}</p>'

    analisi_sito_narr = ""
    if narr.get("analisi_sito"):
        analisi_sito_narr = f'<p class="narrative">{_escape_html(narr["analisi_sito"])}</p>'

    risorsa_solare_narr = ""
    if narr.get("risorsa_solare"):
        risorsa_solare_narr = f'<p class="narrative">{_escape_html(narr["risorsa_solare"])}</p>'

    dimensionamento_narr = ""
    if narr.get("dimensionamento"):
        dimensionamento_narr = f'<p class="narrative">{_escape_html(narr["dimensionamento"])}</p>'

    conclusioni_html = ""
    if narr.get("conclusioni"):
        conclusioni_html = f'<h2>7. Conclusioni e raccomandazioni</h2><p class="narrative">{_escape_html(narr["conclusioni"])}</p>'

    return f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="utf-8">
<title>Capitolato Tecnico - Impianto Fotovoltaico</title>
<style>
    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; line-height: 1.6; }}
    h1 {{ color: #1a5276; border-bottom: 3px solid #f39c12; padding-bottom: 10px; }}
    h2 {{ color: #2c3e50; margin-top: 30px; border-left: 4px solid #f39c12; padding-left: 12px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
    td {{ padding: 8px 12px; border-bottom: 1px solid #eee; }}
    td:first-child {{ font-weight: 600; width: 40%; color: #555; }}
    .header {{ text-align: center; margin-bottom: 30px; }}
    .date {{ text-align: right; color: #777; font-size: 0.9em; }}
    .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.85em; color: #777; text-align: center; }}
    ul {{ padding-left: 20px; }}
    li {{ margin-bottom: 5px; }}
    .norms {{ background: #f8f9fa; padding: 15px; border-radius: 5px; }}
    .narrative {{ background: #f0f7ff; padding: 14px 18px; border-left: 4px solid #2980b9; border-radius: 0 5px 5px 0; margin: 12px 0; font-style: italic; color: #2c3e50; }}
</style>
</head>
<body>
<div class="header">
    <h1>Capitolato Tecnico &mdash; Impianto Fotovoltaico</h1>
    <p class="date">Data: {today}</p>
</div>

{premessa_html}

<h2>1. Dati del sito</h2>
{analisi_sito_narr}
<table>
    <tr><td>Indirizzo</td><td>{design.site.address}</td></tr>
    <tr><td>Coordinate</td><td>{design.site.latitude:.5f}&deg;N, {design.site.longitude:.5f}&deg;E</td></tr>
    <tr><td>Comune</td><td>{design.site.municipality} ({design.site.province})</td></tr>
    <tr><td>Regione</td><td>{design.site.region}</td></tr>
    <tr><td>Zona climatica</td><td>{design.site.climate_zone}</td></tr>
    <tr><td>Zona sismica</td><td>{design.site.seismic_zone}</td></tr>
</table>

<h2>2. Analisi solare</h2>
{risorsa_solare_narr}
<table>
    <tr><td>Irraggiamento annuo</td><td>{design.solar_data.annual_irradiation} kWh/m&sup2;/anno</td></tr>
    <tr><td>Inclinazione ottimale</td><td>{design.solar_data.optimal_tilt}&deg;</td></tr>
    <tr><td>Azimut ottimale</td><td>{design.solar_data.optimal_azimuth}&deg;</td></tr>
    <tr><td>Producibilit&agrave; specifica</td><td>{design.solar_data.annual_production_per_kwp} kWh/kWp/anno</td></tr>
</table>
{"<h3>Irraggiamento mensile</h3><table>" + monthly_rows + "</table>" if monthly_rows else ""}

<h2>3. Dimensionamento impianto</h2>
{dimensionamento_narr}
<table>
    <tr><td>Potenza nominale</td><td>{design.system_size_kwp} kWp</td></tr>
    <tr><td>Numero moduli</td><td>{design.num_panels}</td></tr>
    {module_html}
    {inverter_html}
    <tr><td>Produzione annua stimata</td><td>{design.estimated_production_kwh:.0f} kWh</td></tr>
    <tr><td>Autoconsumo stimato</td><td>{design.self_consumption_rate}%</td></tr>
    <tr><td>Performance Ratio</td><td>{design.performance_ratio}</td></tr>
</table>

{economics_html}
{notes_html}

<h2>{"6" if design.notes else "5"}. Riferimenti normativi</h2>
<div class="norms">
<ul>
    <li>CEI 0-21 &mdash; Regola tecnica di connessione utenti attivi BT</li>
    <li>CEI 0-16 &mdash; Regola tecnica di connessione utenti attivi MT</li>
    <li>D.Lgs. 199/2021 &mdash; Attuazione direttiva RED II</li>
    <li>DM 14/01/2008 &mdash; Norme tecniche costruzioni (NTC)</li>
    <li>D.L. 63/2013 &mdash; Detrazioni fiscali per ristrutturazione edilizia</li>
    <li>Delibera ARERA 03/2020 &mdash; Regolazione SSP (Scambio Sul Posto)</li>
</ul>
</div>

{conclusioni_html}

<div class="footer">
    <p>Documento generato con SolarSpec v0.1.0</p>
</div>
</body>
</html>"""


def _generate_pdf(
    design: SystemDesign, output_path: str, narrative: dict[str, str] | None = None
) -> str:
    """Generate a PDF technical specification using WeasyPrint."""
    try:
        from weasyprint import HTML
    except ImportError:
        raise ImportError("Installa weasyprint: pip install weasyprint")

    html_content = _build_html(design, narrative=narrative)
    HTML(string=html_content).write_pdf(output_path)
    return output_path
